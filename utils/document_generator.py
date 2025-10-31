from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import re


def sanitize_text(text):
    """Clean text for DOCX compatibility - removes invalid XML characters"""
    if not text or not isinstance(text, str):
        return ""
    
    # Remove control characters and invalid XML chars (except newline, carriage return, tab)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Remove null bytes and BOM
    text = text.replace('\x00', '').replace('\ufeff', '')
    
    # Ensure text is string and not None
    return str(text).strip() if text else ""


def extract_user_details(resume_text):
    """Extract user name, email, and phone from resume text"""
    resume_text = sanitize_text(resume_text)
    lines = [line.strip() for line in resume_text.split('\n') if line.strip()]
    
    # Extract name
    name = "Your Name"
    for line in lines[:10]:  # Check first 10 lines only
        clean_line = sanitize_text(line)
        if 1 < len(clean_line.split()) <= 4 and all(w[0].isupper() for w in clean_line.split() if w and w[0].isalpha()) and '@' not in clean_line:
            name = clean_line
            break
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume_text)
    email = sanitize_text(email_match.group(0)) if email_match else "your.email@example.com"
    
    # Extract phone
    phone_match = re.search(r'(\+91[-\s]?|0)?[6-9]\d{9}', resume_text)
    phone = sanitize_text(phone_match.group(0)) if phone_match else "+91-xxxxxxxxxx"
    
    return {
        "name": sanitize_text(name),
        "email": email,
        "phone": phone
    }


def create_resume_docx(resume_text, job_keywords, score_data):
    """Create optimized resume DOCX with sanitized content"""
    try:
        resume_text = sanitize_text(resume_text)
        user = extract_user_details(resume_text)
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Add header with name
        p = doc.add_paragraph()
        run = p.add_run(user['name'])
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact_text = f"{user['email']} | {user['phone']}"
        doc.add_paragraph(sanitize_text(contact_text))
        
        # Add skills section
        doc.add_heading('SKILLS', level=1)
        matched = [sanitize_text(s) for s in score_data.get('matched_keywords', [])]
        missing = [sanitize_text(s) for s in score_data.get('missing_keywords', [])]
        all_skills = sorted(set(matched) | set(missing))
        all_skills = [s for s in all_skills if s]  # Remove empty strings
        
        if all_skills:
            doc.add_paragraph(', '.join(all_skills[:30]))  # Limit to 30 skills
        
        # Add original resume content
        doc.add_heading('ORIGINAL RESUME', level=1)
        lines = resume_text.splitlines()[:50]  # Limit to 50 lines
        for line in lines:
            clean_line = sanitize_text(line)
            if clean_line:
                try:
                    doc.add_paragraph(clean_line)
                except Exception:
                    # Skip problematic lines
                    continue
        
        # Add page break
        doc.add_page_break()
        
        # Add ATS improvement suggestions
        doc.add_heading('ATS Missing Skills to Add', level=1)
        if missing:
            doc.add_paragraph("Add these skills to boost your ATS score:")
            for skill in missing[:15]:  # Limit to 15 missing skills
                if skill:
                    try:
                        doc.add_paragraph(skill, style='List Bullet')
                    except Exception:
                        doc.add_paragraph(f"• {skill}")
        else:
            doc.add_paragraph("Great! Your resume matches all required skills.")
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
        
    except Exception as e:
        # Fallback: create minimal document
        doc = Document()
        doc.add_heading('Resume Generation Error', level=1)
        doc.add_paragraph(f"Unable to generate full resume. Error: {str(e)[:100]}")
        doc.add_paragraph(f"ATS Score: {score_data.get('ats_score', 0)}%")
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes


def create_cover_letter_docx(resume_text, job_keywords, score_data):
    """Create cover letter DOCX with sanitized content"""
    try:
        resume_text = sanitize_text(resume_text)
        user = extract_user_details(resume_text)
        
        doc = Document()
        
        # Add user details
        doc.add_paragraph(user['name'])
        contact_line = f"{user['email']} | {user['phone']}"
        doc.add_paragraph(sanitize_text(contact_line))
        doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        
        # Add company and position
        company = sanitize_text(job_keywords.get('company', 'Hiring Manager'))
        job_title = sanitize_text(job_keywords.get('job_title', 'the position'))
        
        doc.add_paragraph(company)
        doc.add_paragraph(f"Re: {job_title}")
        doc.add_paragraph("")
        doc.add_paragraph("Dear Hiring Manager,")
        doc.add_paragraph("")
        
        # Add body content
        matched_skills = [sanitize_text(s) for s in score_data.get('matched_keywords', [])[:3]]
        matched_skills = [s for s in matched_skills if s]
        skills_str = ', '.join(matched_skills) if matched_skills else 'relevant technical skills'
        
        body_text = f"I am writing to express my strong interest in the {job_title} position at {company}. With expertise in {skills_str}, I am confident in my ability to contribute effectively to your team."
        doc.add_paragraph(sanitize_text(body_text))
        doc.add_paragraph("")
        
        ats_text = f"My resume demonstrates a {score_data.get('ats_score', 0)}% match with your job requirements, reflecting my alignment with the role's technical demands."
        doc.add_paragraph(sanitize_text(ats_text))
        doc.add_paragraph("")
        
        doc.add_paragraph("I look forward to the opportunity to discuss how my background and skills align with your needs.")
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(user['name'])
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
        
    except Exception as e:
        # Fallback document
        doc = Document()
        doc.add_heading('Cover Letter', level=1)
        doc.add_paragraph("Dear Hiring Manager,")
        doc.add_paragraph(f"I am interested in the position at your company.")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("Your Name")
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes


def create_ats_report_docx(score_data, job_keywords):
    """Create ATS analysis report DOCX with sanitized content"""
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('ATS ANALYSIS REPORT', level=1)
        doc.add_paragraph("")
        
        # Add score
        doc.add_heading('Overall ATS Score', level=2)
        doc.add_paragraph(f"Score: {score_data.get('ats_score', 0)}%")
        doc.add_paragraph(f"Total Keywords Analyzed: {score_data.get('total_keywords', 0)}")
        doc.add_paragraph(f"Matched: {score_data.get('matched_count', 0)}")
        doc.add_paragraph(f"Missing: {score_data.get('missing_count', 0)}")
        doc.add_paragraph("")
        
        # Add matched keywords
        matched = [sanitize_text(s) for s in score_data.get('matched_keywords', [])]
        matched = [s for s in matched if s]
        
        if matched:
            doc.add_heading('✅ Matched Skills', level=2)
            matched_text = ', '.join(matched[:20])  # Limit to 20
            doc.add_paragraph(sanitize_text(matched_text))
            doc.add_paragraph("")
        
        # Add missing keywords
        missing = [sanitize_text(s) for s in score_data.get('missing_keywords', [])]
        missing = [s for s in missing if s]
        
        if missing:
            doc.add_heading('❌ Missing Skills to Add', level=2)
            missing_text = ', '.join(missing[:20])  # Limit to 20
            doc.add_paragraph(sanitize_text(missing_text))
            doc.add_paragraph("")
        else:
            doc.add_heading('✅ All Skills Matched!', level=2)
            doc.add_paragraph("Your resume contains all the required skills from the job description.")
            doc.add_paragraph("")
        
        # Add job details
        doc.add_heading('Job Details', level=2)
        job_title = sanitize_text(job_keywords.get('job_title', 'N/A'))
        company = sanitize_text(job_keywords.get('company', 'N/A'))
        
        doc.add_paragraph(f"Position: {job_title}")
        doc.add_paragraph(f"Company: {company}")
        doc.add_paragraph(f"Must-have Skills: {len(job_keywords.get('must_have', []))}")
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
        
    except Exception as e:
        # Fallback document
        doc = Document()
        doc.add_heading('ATS Report', level=1)
        doc.add_paragraph(f"ATS Score: {score_data.get('ats_score', 0)}%")
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
