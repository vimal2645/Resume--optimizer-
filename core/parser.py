import PyPDF2
import docx
import re
import json
import io

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except:
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        return text.strip()
    except:
        return ""

SKILLS = [
    "python", "java", "javascript", "react", "node", "express", "django", "flask",
    "sql", "mysql", "postgresql", "mongodb", "docker", "kubernetes", "git",
    "machine learning", "deep learning", "tensorflow", "pytorch", "nlp", "ai",
    "data analysis", "data analytics", "pandas", "numpy", "matplotlib", "seaborn",
    "aws", "azure", "gcp", "rest api", "html", "css", "excel", "ms-excel",
    "power bi", "tableau", "artificial intelligence", "aiml", "bioinformatics", "mlops",
    "scikit-learn", "keras", "apache", "spark", "hadoop", "scala", "r programming",
    "jenkins", "gitlab", "github", "bitbucket", "api", "json", "xml"
]

def extract_skills(text: str) -> list:
    text_lower = text.lower()
    found = set()
    for skill in SKILLS:
        if skill in text_lower:
            found.add(skill)
    return sorted(list(found))

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith('.docx'):
        text = extract_text_from_docx(file_bytes)
    else:
        return {'text': '', 'skills': []}
    skills = extract_skills(text)
    return {'text': text, 'skills': skills}

def parse_job_description_fallback(job_text: str) -> dict:
    """Fallback: extract skills without LLM"""
    skills_found = []
    text_lower = job_text.lower()
    
    for skill in SKILLS:
        if skill in text_lower:
            skills_found.append(skill)
    
    # Try to find job title
    title_match = re.search(r'(job title|position|role):\s*([a-zA-Z\s]+)', job_text, re.IGNORECASE)
    title = title_match.group(2).strip() if title_match else 'Position'
    
    # Try to find company
    company_match = re.search(r'(company|employer):\s*([a-zA-Z\s]+)', job_text, re.IGNORECASE)
    company = company_match.group(2).strip() if company_match else 'Company'
    
    return {
        "must_have": list(set(skills_found[:15])),
        "nice_to_have": [],
        "experience": "Entry",
        "job_title": title,
        "company": company
    }

def parse_job_description(job_text: str) -> dict:
    from utils.llm import get_llm_response
    
    prompt = f"""Extract job requirements from this text and return ONLY valid JSON:

{job_text}

Return this exact format (no markdown, no extra text):
{{"must_have": ["skill1", "skill2"], "nice_to_have": ["skill3"], "experience": "Entry", "job_title": "Title", "company": "Company"}}"""
    
    try:
        response = get_llm_response(prompt)
        
        if not response or response == "{}":
            return parse_job_description_fallback(job_text)
        
        cleaned = response.strip()
        
        # Remove markdown code blocks and extract JSON if present
        m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", cleaned, re.S)
        if m:
            cleaned = m.group(1).strip()
        else:
            # Strip any leading/trailing fenced backticks if present
            if cleaned.startswith("```"):
                cleaned = cleaned.lstrip("`").strip()
            if cleaned.endswith("```"):
                cleaned = cleaned.rstrip("`").strip()
        
        data = json.loads(cleaned)
        
        if not isinstance(data.get('must_have'), list):
            data['must_have'] = []
        if not isinstance(data.get('nice_to_have'), list):
            data['nice_to_have'] = []
        if not data.get('job_title'):
            data['job_title'] = 'Position'
        if not data.get('company'):
            data['company'] = 'Company'
        if not data.get('experience'):
            data['experience'] = 'Entry'
        
        return data
    
    except:
        return parse_job_description_fallback(job_text)
